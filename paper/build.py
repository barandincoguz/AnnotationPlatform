"""Build the IISEC paper from content.md into the conference template."""
import re
import shutil
import subprocess
from pathlib import Path

import docx
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "iisec_template.docx"
CONTENT = Path(__file__).parent / "content.md"
OUT = Path(__file__).parent / "out"
SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Style names verified against the template in Step 1. Exact strings.
STYLE_TITLE = "paper title"
STYLE_AUTHOR = "Author"
STYLE_ABSTRACT = "Abstract"
STYLE_KEYWORDS = "Keywords"
STYLE_H1 = "Heading 1"
STYLE_H2 = "Heading 2"
STYLE_H5 = "Heading 5"
STYLE_BODY = "Body Text"
STYLE_REFS = "references"
STYLE_FIGCAP = "figure caption"
STYLE_TABHEAD = "table head"
STYLE_TABCOLHEAD = "table col head"
STYLE_TABCOPY = "table copy"


def parse_content(path: Path) -> list[tuple[str, str, str]]:
    """Return [(kind, label, text)] where kind is TITLE/ABSTRACT/KEYWORDS/H1/H2/REFERENCES."""
    blocks: list[tuple[str, str, str]] = []
    # All markers are declared here once, including FIGURE and TABLE which Tasks 3 and 4 use.
    # Declaring them up front means neither task has to edit this pattern.
    marker = re.compile(
        r"^<!--\s*(TITLE|AUTHORS|ABSTRACT|KEYWORDS|REFERENCES|FIGURE|TABLE|H1|H2):?\s*(.*?)\s*-->$"
    )
    kind, label, buf = None, "", []

    def flush():
        if kind is None:
            return
        body = "\n".join(buf).strip()
        blocks.append((kind, label, body))

    for line in path.read_text(encoding="utf-8").splitlines():
        m = marker.match(line.strip())
        if m:
            flush()
            kind, label, buf = m.group(1), m.group(2), []
        else:
            buf.append(line)
    flush()
    return blocks


def strip_template_content(doc: docx.Document) -> list:
    """Delete the template's boilerplate while PRESERVING the paragraphs whose pPr carries a
    sectPr. Those paragraphs encode the column layout (1-col title, 3-col author, 2-col body);
    deleting them collapses the document into a single section and the title typesets inside a
    narrow body column. Their runs are emptied so no template text survives.

    Returns the surviving structural anchors in document order:
      [0] paragraph ending the 1-column title section
      [1] paragraph ending the first 3-column author section
      [2] paragraph ending the duplicate 3-column section
      [3] paragraph ending the 2-column body section
      [4] the body-level sectPr element
    """
    body = doc.element.body
    anchors = []
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            anchors.append(child)
            continue
        if child.find(f"{NS}pPr/{NS}sectPr") is not None:
            for run in child.findall(f"{NS}r"):
                child.remove(run)
            anchors.append(child)
            continue
        body.remove(child)
    if len(anchors) != 5:
        raise RuntimeError(f"expected 5 structural anchors, found {len(anchors)} — template changed")
    return anchors


def para_before(doc: docx.Document, anchor, style: str, text: str = ""):
    """Create a paragraph and move it immediately before `anchor`, so it lands in the section
    that `anchor` terminates. python-docx can only append, so we append then relocate."""
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    anchor.addprevious(p._element)
    return p


def insert_figure(doc: docx.Document, anchor, png_path: Path, caption: str, width_pt: float = 252):
    """Insert a figure (as an image paragraph) followed by its caption, both immediately
    before `anchor`. The caption uses STYLE_FIGCAP so it auto-numbers as "Fig. 1."."""
    p = doc.add_paragraph()
    p.add_run().add_picture(str(png_path), width=Pt(width_pt))
    anchor.addprevious(p._element)
    para_before(doc, anchor, STYLE_FIGCAP, caption)


# The body section is 2 columns; a table must fit inside one of them. python-docx's
# add_table() defaults each column to a full single-column page width (~253pt), which
# overflows the ~243pt IEEE column, so widths must be fixed explicitly.
TABLE_WIDTH_PT = 241


_NUMERIC_TOKEN = re.compile(r"[0-9.,/%]+")


def _token_lengths(cell: str) -> tuple[int, int]:
    """Longest numeric-token length and longest word-token length in `cell` (split on
    whitespace and hyphens). Distinguished because a wrapped numeric value like "0.7625"
    reads as a different, wrong number, while a wrapped header word like "Precision" is
    merely a stylistic compromise a dense table can afford."""
    max_num = max_word = 0
    for tok in re.split(r"[\s-]+", cell.strip()):
        if not tok:
            continue
        if _NUMERIC_TOKEN.fullmatch(tok):
            max_num = max(max_num, len(tok))
        else:
            max_word = max(max_word, len(tok))
    return max_num, max_word


def _col_widths(header: list[str], rows: list[list[str]], total_pt: float = TABLE_WIDTH_PT,
                 char_pt: float = 4.3, pad_pt: float = 11.0) -> list[float]:
    """Column widths (points) summing to total_pt.

    Each column gets a floor wide enough to hold its longest numeric token whole and its
    longest word token across at most two wrapped lines, rather than fragmenting character-
    by-character (e.g. a bare length-proportional split gave "Precision" only ~23pt, which
    rendered as "Pr"/"eci"/"sio"/"n", and "0.789" as "0.78"/"9"). Remaining budget is
    distributed proportionally by each column's longest cell text, so a column holding long
    free-form text gets more room to wrap across fewer lines."""
    ncols = len(header)
    max_num = [0] * ncols
    max_word = [0] * ncols
    total_len = [len(h) for h in header]
    for cells in (header, *rows):
        for i, cell in enumerate(cells):
            total_len[i] = max(total_len[i], len(cell))
            n, w = _token_lengths(cell)
            max_num[i] = max(max_num[i], n)
            max_word[i] = max(max_word[i], w)
    floor = []
    for i in range(ncols):
        full = max_num[i] * char_pt + pad_pt if max_num[i] else 0
        half = ((max_word[i] + 1) // 2) * char_pt + pad_pt if max_word[i] else 0
        floor.append(max(full, half, pad_pt))
    if sum(floor) >= total_pt:
        scale = total_pt / sum(floor)
        return [f * scale for f in floor]
    remaining = total_pt - sum(floor)
    total_w = sum(total_len)
    widths = [f + remaining * w / total_w for f, w in zip(floor, total_len)]
    scale = total_pt / sum(widths)
    return [w * scale for w in widths]


def insert_table(doc: docx.Document, anchor, caption: str, header: list[str], rows: list[list[str]]):
    """Insert a table caption (auto-numbered "TABLE I.", etc.) followed by a table, both
    immediately before `anchor`. The template has no 'Table Grid' style, so 'Normal Table'
    is used and the cell paragraphs get the template's own header/copy styles."""
    para_before(doc, anchor, STYLE_TABHEAD, caption)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Normal Table"
    t.autofit = False
    widths = [Pt(w) for w in _col_widths(header, rows)]
    for col, w in zip(t.columns, widths):
        col.width = w  # fixes column width for rows added below via add_row()
    for cell, name, w in zip(t.rows[0].cells, header, widths):
        cell.width = w  # the header row already existed, so its cells need fixing directly
        cell.paragraphs[0].style = doc.styles[STYLE_TABCOLHEAD]
        cell.paragraphs[0].add_run(name)
    for row in rows:
        for cell, value in zip(t.add_row().cells, row):
            cell.paragraphs[0].style = doc.styles[STYLE_TABCOPY]
            cell.paragraphs[0].add_run(value)
    anchor.addprevious(t._element)
    return t


def render(blocks, doc: docx.Document, anchors: list) -> None:
    title_anchor, author_anchor, body_anchor = anchors[0], anchors[1], anchors[3]
    for kind, label, text in blocks:
        if kind == "TITLE":
            para_before(doc, title_anchor, STYLE_TITLE, text)
        elif kind == "AUTHORS":
            for line in [l for l in text.splitlines() if l.strip()]:
                para_before(doc, author_anchor, STYLE_AUTHOR, line.strip())
        elif kind == "ABSTRACT":
            para_before(doc, body_anchor, STYLE_ABSTRACT, f"Abstract—{text}")
        elif kind == "KEYWORDS":
            para_before(doc, body_anchor, STYLE_KEYWORDS, f"Keywords—{text}")
        elif kind in ("H1", "H2"):
            para_before(doc, body_anchor, STYLE_H1 if kind == "H1" else STYLE_H2, label)
            for para in [p for p in text.split("\n\n") if p.strip()]:
                para_before(doc, body_anchor, STYLE_BODY, para.strip())
        elif kind == "REFERENCES":
            para_before(doc, body_anchor, STYLE_H5, "References")
            for entry in [l for l in text.splitlines() if l.strip()]:
                para_before(doc, body_anchor, STYLE_REFS, entry.strip())
        elif kind == "FIGURE":
            insert_figure(doc, body_anchor, OUT / "fig1.png", text)
        elif kind == "TABLE":
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            caption, row_lines = lines[0], lines[1:]
            stray = [l for l in row_lines if not l.startswith("|")]
            if stray:
                raise ValueError(
                    f"TABLE block {caption!r} has non-row content: {stray!r}. "
                    "Prose after a table is silently parsed as table rows — put it under "
                    "the next H1/H2 marker instead."
                )
            grid = [[c.strip() for c in l.strip("|").split("|")] for l in row_lines]
            insert_table(doc, body_anchor, caption, grid[0], grid[1:])
        else:
            raise ValueError(f"unknown content marker: {kind!r}")


def docx_to_pdf(docx_path: Path, outdir: Path) -> Path:
    """Convert with an isolated LibreOffice profile so a running GUI instance cannot block us."""
    subprocess.run(
        [SOFFICE, "-env:UserInstallation=file:///tmp/lo_iisec_profile",
         "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
        check=True, capture_output=True,
    )
    pdf = outdir / (docx_path.stem + ".pdf")
    if not pdf.exists():
        raise RuntimeError(f"LibreOffice produced no PDF at {pdf}")
    return pdf


def build() -> tuple[Path, Path]:
    OUT.mkdir(exist_ok=True)
    work = OUT / "paper.docx"
    shutil.copy(TEMPLATE, work)
    doc = docx.Document(work)
    anchors = strip_template_content(doc)
    render(parse_content(CONTENT), doc, anchors)
    doc.save(work)
    pdf = docx_to_pdf(work, OUT)
    return work, pdf


if __name__ == "__main__":
    d, p = build()
    print(f"docx: {d}\npdf:  {p}")
