"""Verification gates for the IISEC paper. Exit 0 = pass, 1 = fail."""
import re
import sys
from pathlib import Path

import docx
import fitz

OUT = Path(__file__).parent / "out"
DOCX = OUT / "paper.docx"
PDF = OUT / "paper.pdf"
MAX_PAGES = 6

# Every numeric token permitted in the paper, from spec section 5.
ALLOWED_NUMBERS = {
    # corpus
    "17923", "17,923", "1437", "1,437", "1413", "1,413", "6840", "6,840", "14", "4.84",
    "577", "29.1", "6", "200", "0",
    # model
    "3.5", "9", "4", "8", "16", "20.0", "0.0", "1", "42", "2.5", "1.0", "1536", "256",
    "12288", "12,288", "8253", "8,253", "494", "4278", "4,278", "1003", "1,003", "0.94", "0.31.0",
    # model architecture (mlx-lm loader description, spec 5.2)
    "32", "27", "4096", "4,096", "248320", "248,320", "64",
    # splits
    "500", "394", "50", "17423", "17,423", "38",
    # results
    "0.789", "0.861", "0.728", "13", "26.0", "0.805", "0.8525", "0.7625", "47", "47.0", "100",
    # routing
    "1294", "1,294", "342", "211", "738", "3", "949", "26.4", "16.3", "57.0", "0.2", "73.3", "100",
    # review load counting QUARANTINE on the review side: 211 + 738 + 3 (spec 5.5)
    "952",
    # runtime
    "7.63", "6.21", "13.26", "17.29", "24.62", "36.83", "9872", "9,872", "10964", "10,964",
    "2", "45", "11", "12", "3.05",
    # development configuration and checkpoint selection (spec 5.2)
    "75", "150", "550", "3399", "3,399",
    # windowing coverage and reported confidence level (spec 5.2, canonical since 9bbfa82)
    "99", "95",
    # canonical schema example (spec 5.1)
    "3065", "298",
    # years, citation numbers and ordinals are handled separately
}

# Terms that must never appear, mapped to the reason.
FORBIDDEN = {
    "harness": "quality harness is out of scope (spec 3)",
    "jaccard": "inter-annotator Jaccard is out of scope (spec 3)",
    "kappa": "chance-corrected agreement is out of scope (spec 3)",
    "krippendorff": "out of scope (spec 3)",
    "gamif": "gamification is out of scope (spec 3)",
    "review_kept": "out of scope (spec 3)",
    "outbox": "CDC mirror is out of scope (spec 3)",
    "19.58": "pre-policy comparison is excluded (spec 3)",
    "1,180": "pre-policy review load is excluded (spec 3)",
    "1180": "pre-policy review load is excluded (spec 3)",
    "1,087": "pre-policy bucket count is excluded (spec 3)",
    "1087": "pre-policy bucket count is excluded (spec 3)",
    # residual template boilerplate
    "use style": "template guidance text not removed",
    "Heading 1)": "template guidance text not removed",
    "IEEE conference templates contain guidance text": "template warning not removed",
    "Identify applicable funding agency": "sponsor placeholder not removed",
    "G. Eason": "template example reference not removed",
    "K. Elissa": "template example reference not removed",
}

# Entries matched as regexes rather than plain substrings, for terms whose bare
# digits collide with legitimate content (e.g. a reference page range 406-413).
# A plain "413" substring would also match reference 11's page range and the canonical
# 1,413 completed-document count, so the policy mention is targeted instead of the digits.
FORBIDDEN_PATTERNS = {
    r"\b213\s*/\s*413\b": "VUK 213/413 policy is excluded by user decision (spec 3)",
    r"\b(?:VUK|Article|article|madde)\s*413\b": "VUK 413 provision is excluded (spec 3)",
}

# Numbers that are not merely absent from the allowlist but must never be added to it:
# the pre-policy bucket counts and transitions excluded by spec section 3. They are
# already blocked by the allowlist; this set only corrects the guidance in the message.
NEVER_ADD = {"111", "93", "231", "118", "1087", "1,087", "1180", "1,180", "19.58"}


def extract_text(path: Path, claims_only: bool = False) -> str:
    """Full document text, or only the claim-bearing text before the References heading.

    Number provenance applies to claims, not to citation metadata: reference entries
    legitimately carry volumes, issues, page ranges, arXiv IDs and DOIs that are not
    facts about this work and do not belong in the canonical facts table.
    """
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        if claims_only and p.text.strip().lower() == "references":
            break
        parts.append(p.text)
    for t in d.tables:
        # Boundary marker: without it, one table's trailing cells and the next table's
        # leading cells sit back-to-back with no gap, so check_bucket_figures' 40-character
        # window bleeds an unrelated table's numbers into a bucket name in the next table.
        parts.append("=" * 80)
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def check_forbidden(text: str) -> list[str]:
    low = text.lower()
    out = [f"FORBIDDEN {term!r}: {why}" for term, why in FORBIDDEN.items()
           if term.lower() in low]
    out += [f"FORBIDDEN /{pat}/: {why}" for pat, why in FORBIDDEN_PATTERNS.items()
            if re.search(pat, text)]
    return out


def check_numbers(text: str) -> list[str]:
    """Every number outside a citation bracket must be in the allowlist."""
    stripped = re.sub(r"\[\d{1,2}\](?:\s*,\s*\[\d{1,2}\])*", " ", text)      # citations [1], [2]
    stripped = re.sub(r"\b(19|20|26)\d{2}\b", " ", stripped)                  # years
    stripped = re.sub(r"\b[IVX]+\.", " ", stripped)                           # roman headings
    stripped = re.sub(r"\b(?=[0-9a-f]*[a-f])[0-9a-f]{6,}\b", " ", stripped)   # hex hashes ONLY
    stripped = re.sub(r"\b\d+\.\d+\.\d+\b", " ", stripped)                    # versions 0.31.0
    stripped = re.sub(r"\d+\.?\d*\s*[eE]\s*[-−]?\d+", " ", stripped)     # 2.5e-5
    stripped = re.sub(r"\b[A-Za-z_]\w*(?:[._]\w+)+", " ", stripped)           # text_config, Qwen3.5
    stripped = re.sub(r"\bp(?:50|90|95|99)\b", " ", stripped)      # percentile labels
    bad = []
    for tok in re.findall(r"\d(?:[\d,]*\d)?(?:\.\d+)?", stripped):
        if tok not in ALLOWED_NUMBERS:
            if tok in NEVER_ADD:
                bad.append(f"FORBIDDEN NUMBER {tok!r} — pre-policy figure, must NEVER "
                           f"appear and must NEVER be added to ALLOWED_NUMBERS")
            else:
                bad.append(f"UNKNOWN NUMBER {tok!r} — add to spec section 5 first")
    return sorted(set(bad))


# Canonical routing figures, spec section 5.5.
BUCKET_CANON = {"GREEN": 342, "YELLOW": 211, "RED": 738, "QUARANTINE": 3}

# Every value that may legitimately appear within forty characters of a bucket name:
# the routing counts (spec 5.5), the four shares and the batch totals, plus the rates and
# scores that could plausibly be discussed alongside a bucket. Anything else near a bucket
# name is a figure error. Compared as floats so 57 and 57.0 are the same value.
ROUTING_FIGURES = set(BUCKET_CANON.values()) | {1294, 949}
NEAR_BUCKET_OK = ({float(v) for v in ROUTING_FIGURES}
                  | {26.4, 16.3, 57.0, 0.2, 73.3}          # shares
                  | {4.84, 0.789, 0.805, 0.861, 0.728})    # rates and scores


def check_bucket_figures(text: str) -> list[str]:
    """Flag any number overlapping a forty-character window around a bucket name whose value
    is not in NEAR_BUCKET_OK.

    Number spans are located once and selected by overlap rather than by slicing the text,
    so a number straddling the window boundary is evaluated whole while a number starting at
    or after the boundary is excluded. Slicing plus outward snapping got both cases wrong.

    Years and section, figure, table, step and task references are removed first. Decimals
    are deliberately NOT removed: scrubbing them to avoid flagging rates also hid
    decimal-formatted wrong counts such as 700.0 against a canonical 738.
    """
    errs = []
    scrubbed = re.sub(r"\b(19|20|26)\d{2}\b", " ", text)
    scrubbed = re.sub(r"\b(?:section|sec\.|fig\.|figure|table|step|task)\s*[IVX\d]+",
                      " ", scrubbed, flags=re.I)
    scrubbed = re.sub(r"(\d[\d,]*(?:\.\d+)?)\s*%", lambda m: " " * len(m.group(0)), scrubbed)
    spans = [(m.start(), m.end(), float(m.group(0).replace(",", "")))
             for m in re.finditer(r"\d[\d,]*(?:\.\d+)?", scrubbed)]
    for name, canon in BUCKET_CANON.items():
        for m in re.finditer(rf"\b{name}\b", scrubbed, re.I):
            lo = max(0, m.start() - 40)
            hi = min(len(scrubbed), m.end() + 40)
            for s0, s1, value in spans:
                if s1 > lo and s0 < hi and value not in NEAR_BUCKET_OK:
                    errs.append(
                        f"BUCKET FIGURE: {name} appears near {value:g}, which is not a "
                        f"legitimate figure (canonical count is {canon})"
                    )
    return sorted(set(errs))


def check_bucket_table_rows() -> list[str]:
    """In any table row whose first cell names a bucket, every integer cell must be a
    legitimate routing figure and the bucket's canonical count must be among them.

    Inspects all cells rather than stopping at the first integer, because a decoy numeric
    cell before the count column otherwise masks a wrong count. Trailing footnote markers
    and whitespace are stripped before matching, because a cell reading '50*' otherwise
    fails to parse and the row goes unchecked."""
    errs = []
    doc = docx.Document(DOCX)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            for name, canon in BUCKET_CANON.items():
                if not re.search(rf"\b{name}\b", cells[0], re.I):
                    continue
                ints = []
                for cell in cells[1:]:
                    clean = re.sub(r"[*†‡\s]+$", "", cell)
                    if re.fullmatch(r"\d[\d,]*", clean):
                        ints.append(int(clean.replace(",", "")))
                foreign = [v for v in ints if v not in ROUTING_FIGURES]
                if foreign:
                    errs.append(
                        f"TABLE {ti + 1} row {ri + 1}: bucket {name} row contains "
                        f"{foreign}, not legitimate routing figures "
                        f"(canonical count is {canon})"
                    )
                elif ints and canon not in ints:
                    errs.append(
                        f"TABLE {ti + 1} row {ri + 1}: bucket {name} row states {ints} "
                        f"but not its canonical count {canon}"
                    )
    return errs


def check_table_sums(text: str) -> list[str]:
    """Any table row whose first cell is 'Total' must equal the sum of the numeric rows
    above it, per column. Operates on the rendered document, not on constants."""
    errs = []
    doc = docx.Document(DOCX)
    for ti, table in enumerate(doc.tables):
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        totals = [i for i, r in enumerate(rows) if r and r[0].lower().startswith("total")]
        for tr in totals:
            for col in range(1, len(rows[tr])):
                stated = _as_number(rows[tr][col])
                if stated is None:
                    continue
                above = [_as_number(r[col]) for r in rows[1:tr] if col < len(r)]
                above = [v for v in above if v is not None]
                if _is_count_column(rows, tr, col):
                    tol = 0.0
                else:
                    tol = min(0.05 * len(above) + 0.001, 0.25)
                if above and abs(sum(above) - stated) > tol:
                    errs.append(
                        f"TABLE {ti + 1} column {col}: rows sum to {sum(above)} "
                        f"but the Total row states {stated}"
                    )
    return errs


def _as_number(cell: str):
    m = re.fullmatch(r"(\d(?:[\d,]*\d)?(?:\.\d+)?)\s*%?", cell.strip())
    return float(m.group(1).replace(",", "")) if m else None


def _is_count_column(rows, tr, col) -> bool:
    """True when every cell above the Total row that parses as a number is a plain integer —
    no decimal point, no percent sign. Cells that are not numbers at all (blank, a dash) are
    ignored rather than disqualifying the column, which would silently relax it to the
    tolerance branch. Such a column must sum exactly."""
    cells = [r[col].strip() for r in rows[1:tr] if col < len(r) and r[col].strip()]
    numeric = [c for c in cells if _as_number(c) is not None]
    return bool(numeric) and all(re.fullmatch(r"\d[\d,]*", c) for c in numeric)


def check_checker_constants() -> list[str]:
    """Self-check of this file's own canonical constants. NOT a gate on the document —
    check_bucket_figures and check_table_sums are. Kept so a bad edit to the constants
    above screams instead of silently loosening the gates."""
    errs = []
    if sum(BUCKET_CANON.values()) != 1294:
        errs.append("BUCKET_CANON values do not sum to 1294")
    if BUCKET_CANON["YELLOW"] + BUCKET_CANON["RED"] != 949:
        errs.append("review load is not YELLOW + RED")
    if round(BUCKET_CANON["GREEN"] / 1294 * 100, 1) != 26.4:
        errs.append("26.4% does not follow from GREEN/1294")
    if 394 + 50 + 50 != 494 or 494 + 6 != 500:
        errs.append("canonical split constants are inconsistent")
    return errs


def check_pages() -> list[str]:
    d = fitz.open(PDF)
    if d.page_count > MAX_PAGES:
        return [f"PAGE LIMIT: {d.page_count} pages, maximum is {MAX_PAGES}"]
    return []


def check_no_page_numbers() -> list[str]:
    d = docx.Document(DOCX)
    errs = []
    for i, sec in enumerate(d.sections):
        for footer in (sec.footer, sec.first_page_footer):
            if footer is not None and footer.paragraphs:
                if any(p.text.strip() for p in footer.paragraphs):
                    errs.append(f"section {i}: footer is not empty")
                if footer._element.xpath('.//w:fldSimple | .//w:instrText'):
                    errs.append(f"section {i}: footer contains a field (page number?)")
    return errs


def main() -> int:
    if not DOCX.exists() or not PDF.exists():
        print("FAIL: run paper/build.py first")
        return 1
    full = extract_text(DOCX)
    claims = extract_text(DOCX, claims_only=True)
    problems = (check_forbidden(full) + check_numbers(claims)
                + check_bucket_figures(claims) + check_bucket_table_rows()
                + check_table_sums(claims) + check_checker_constants()
                + check_pages() + check_no_page_numbers())
    if problems:
        print(f"FAIL — {len(problems)} problem(s):")
        for p in problems:
            print("  -", p)
        return 1
    print(f"PASS — {fitz.open(PDF).page_count} pages, all gates clear")
    return 0


if __name__ == "__main__":
    sys.exit(main())
